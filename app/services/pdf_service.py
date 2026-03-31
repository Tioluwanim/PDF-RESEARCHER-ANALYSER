"""
pdf_service.py - Handles PDF upload, storage, and retrieval.
Responsible for saving files to disk and managing document state.
"""

import uuid
import shutil
from pathlib import Path
from datetime import datetime

from app.config import (
    UPLOAD_DIR,
    PROCESSED_DIR,
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB,
    ALLOWED_EXTENSIONS,
)
from app.models.schemas import (
    ProcessedDocument,
    DocumentMetadata,
    DocumentStatus,
    UploadResponse,
    ErrorResponse,
)
from app.utils.logger import get_logger, ServiceLogger

logger = get_logger(__name__)


# ── PDF Service Class ─────────────────────────────────────────────────────────

class PDFService:
    """
    Handles all file-level operations:
    - Validating uploaded PDFs
    - Saving to disk
    - Loading processed document state
    - Listing and deleting documents
    """

    def __init__(self):
        self.upload_dir    = UPLOAD_DIR
        self.processed_dir = PROCESSED_DIR
        logger.info("PDFService initialised")

    # ── Upload ────────────────────────────────────────────────────────────────

    def save_upload(
        self,
        file_bytes: bytes,
        filename: str,
    ) -> tuple[ProcessedDocument, None] | tuple[None, ErrorResponse]:
        """
        Validates and saves an uploaded file to disk.
        Supports: PDF, DOCX, DOC, TXT, XLSX, XLS, CSV.
        Non-PDF files are converted to a text-based PDF wrapper so the
        rest of the pipeline (extraction → embedding → RAG) works unchanged.

        Returns:
            (ProcessedDocument, None) on success.
            (None, ErrorResponse)     on validation failure.
        """
        slog = ServiceLogger("pdf_service")

        suffix = Path(filename).suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            msg = (
                f"Unsupported file type '{suffix}'. "
                f"Supported: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
            )
            slog.warning(msg)
            return None, ErrorResponse(error="Invalid file type", detail=msg)

        file_size = len(file_bytes)
        if file_size == 0:
            return None, ErrorResponse(error="Empty file", detail="Uploaded file is empty.")
        if file_size > MAX_FILE_SIZE_BYTES:
            msg = (
                f"File size {file_size / (1024*1024):.1f} MB exceeds "
                f"the {MAX_FILE_SIZE_MB} MB limit."
            )
            return None, ErrorResponse(error="File too large", detail=msg)

        # ── For non-PDF files: extract text and wrap as a .pdf on disk ────────
        if suffix != ".pdf":
            try:
                text = _extract_text_from_file(file_bytes, suffix, filename)
                if not text.strip():
                    return None, ErrorResponse(
                        error="Empty content",
                        detail=f"Could not extract any text from '{filename}'.",
                    )
                # Wrap as minimal PDF using PyMuPDF
                import fitz
                pdf_doc  = fitz.open()
                page     = pdf_doc.new_page()
                # Insert text with automatic wrapping
                page.insert_textbox(
                    fitz.Rect(50, 50, 562, 792),
                    text,
                    fontsize  = 9,
                    fontname  = "helv",
                    align     = 0,
                )
                # If text overflows one page, add more pages
                if len(text) > 4000:
                    chunks = [text[i:i+4000] for i in range(4000, len(text), 4000)]
                    for chunk in chunks:
                        pg = pdf_doc.new_page()
                        pg.insert_textbox(
                            fitz.Rect(50, 50, 562, 792),
                            chunk, fontsize=9, fontname="helv", align=0,
                        )
                file_bytes = pdf_doc.tobytes()
                pdf_doc.close()
                # Keep original name but save as .pdf
                filename   = Path(filename).stem + ".pdf"
                slog.info(
                    "Converted '%s' (%s) → PDF (%d bytes, %d chars text)",
                    filename, suffix, len(file_bytes), len(text),
                )
            except Exception as e:
                return None, ErrorResponse(
                    error="Conversion failed",
                    detail=f"Could not convert '{filename}' to PDF: {e}",
                )

        # ── Validate PDF magic bytes ──────────────────────────────────────────
        if not file_bytes.startswith(b"%PDF"):
            return None, ErrorResponse(
                error="Invalid PDF",
                detail="File does not appear to be a valid PDF.",
            )

        # ── Save to disk ──────────────────────────────────────────────────────
        doc_id    = str(uuid.uuid4())
        safe_name = self._sanitize_filename(filename)
        dest_path = self.upload_dir / f"{doc_id}_{safe_name}"

        slog = ServiceLogger("pdf_service", doc_id=doc_id)

        # ── Save to disk ──────────────────────────────────────────────────────
        try:
            dest_path.write_bytes(file_bytes)
            slog.info(f"Saved '{filename}' → {dest_path} ({file_size:,} bytes)")
        except OSError as e:
            msg = f"Failed to save file: {e}"
            slog.error(msg)
            return None, ErrorResponse(error="Storage error", detail=msg, doc_id=doc_id)

        # ── Build ProcessedDocument ───────────────────────────────────────────
        doc = ProcessedDocument(
            doc_id    = doc_id,
            filename  = filename,
            file_path = str(dest_path),
            status    = DocumentStatus.UPLOADED,
            metadata  = DocumentMetadata(file_size_bytes=file_size),
        )

        # Persist initial state
        self._save_document_state(doc)

        slog.info(f"Document state created — status={doc.status.value}")
        return doc, None

    # ── Load / Save State ─────────────────────────────────────────────────────

    def load_document(self, doc_id: str) -> ProcessedDocument | None:
        """
        Loads a ProcessedDocument from its JSON state file.

        Args:
            doc_id: The document UUID.

        Returns:
            ProcessedDocument if found, None otherwise.
        """
        state_path = self._state_path(doc_id)
        if not state_path.exists():
            logger.warning(f"[{doc_id}] State file not found: {state_path}")
            return None

        try:
            doc = ProcessedDocument.model_validate_json(state_path.read_text(encoding="utf-8"))
            logger.debug(f"[{doc_id}] Loaded document — status={doc.status.value}")
            return doc
        except Exception as e:
            logger.error(f"[{doc_id}] Failed to load state: {e}", exc_info=True)
            return None

    def save_document(self, doc: ProcessedDocument) -> bool:
        """
        Persists a ProcessedDocument to its JSON state file.
        Updates updated_at timestamp automatically.

        Returns:
            True on success, False on failure.
        """
        doc.updated_at = datetime.utcnow()
        return self._save_document_state(doc)

    def update_status(
        self,
        doc_id: str,
        status: DocumentStatus,
        error_message: str | None = None,
    ) -> bool:
        """
        Convenience method to update only the status of a document.

        Args:
            doc_id:        Document UUID.
            status:        New DocumentStatus value.
            error_message: Optional error detail (set on FAILED status).

        Returns:
            True on success, False if document not found.
        """
        doc = self.load_document(doc_id)
        if not doc:
            return False

        doc.status = status
        if error_message:
            doc.error_message = error_message

        logger.info(f"[{doc_id}] Status → {status.value}")
        return self.save_document(doc)

    # ── List Documents ────────────────────────────────────────────────────────

    def list_documents(self) -> list[dict]:
        """
        Returns a list of all document summaries from state files.
        Used to populate the sidebar in the Streamlit UI.

        Returns:
            List of summary dicts (doc_id, filename, status, pages etc.)
        """
        summaries = []
        for state_file in sorted(
            self.processed_dir.glob("*.json"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,  # most recent first
        ):
            try:
                doc = ProcessedDocument.model_validate_json(
                    state_file.read_text(encoding="utf-8")
                )
                summaries.append(doc.summary())
            except Exception as e:
                logger.warning(f"Could not parse state file {state_file.name}: {e}")

        logger.debug(f"Listed {len(summaries)} documents")
        return summaries

    # ── Delete Document ───────────────────────────────────────────────────────

    def delete_document(self, doc_id: str) -> bool:
        """
        Deletes all files associated with a document:
        - Uploaded PDF
        - State JSON
        - FAISS vector index directory

        Args:
            doc_id: Document UUID.

        Returns:
            True if anything was deleted, False if nothing found.
        """
        slog   = ServiceLogger("pdf_service", doc_id=doc_id)
        doc    = self.load_document(doc_id)
        found  = False

        # Delete uploaded PDF
        if doc and Path(doc.file_path).exists():
            Path(doc.file_path).unlink()
            slog.info("Deleted uploaded PDF")
            found = True

        # Delete vector index
        if doc and doc.vector_index_path:
            idx_path = Path(doc.vector_index_path)
            if idx_path.exists():
                if idx_path.is_dir():
                    shutil.rmtree(idx_path)
                else:
                    idx_path.unlink()
                slog.info("Deleted vector index")

        # Delete state JSON
        state_path = self._state_path(doc_id)
        if state_path.exists():
            state_path.unlink()
            slog.info("Deleted state file")
            found = True

        return found

    # ── Helpers ───────────────────────────────────────────────────────────────

    def get_upload_response(self, doc: ProcessedDocument) -> UploadResponse:
        """Builds an UploadResponse from a ProcessedDocument."""
        return UploadResponse(
            doc_id    = doc.doc_id,
            filename  = doc.filename,
            file_size = doc.metadata.file_size_bytes,
            status    = doc.status,
        )

    def document_exists(self, doc_id: str) -> bool:
        """Returns True if state file exists for this doc_id."""
        return self._state_path(doc_id).exists()

    def is_ready(self, doc_id: str) -> bool:
        """Returns True if document is fully processed and ready for chat."""
        doc = self.load_document(doc_id)
        return doc is not None and doc.status == DocumentStatus.READY

    # ── Private ───────────────────────────────────────────────────────────────

    def _state_path(self, doc_id: str) -> Path:
        """Returns the path to the JSON state file for a document."""
        return self.processed_dir / f"{doc_id}.json"

    def _save_document_state(self, doc: ProcessedDocument) -> bool:
        """Writes the ProcessedDocument as JSON to processed_dir."""
        state_path = self._state_path(doc.doc_id)
        try:
            state_path.write_text(
                doc.model_dump_json(indent=2),
                encoding="utf-8",
            )
            return True
        except OSError as e:
            logger.error(f"[{doc.doc_id}] Failed to save state: {e}", exc_info=True)
            return False

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """
        Removes unsafe characters from a filename.
        Keeps alphanumerics, dots, hyphens and underscores.
        """
        safe = "".join(
            c if c.isalnum() or c in (".", "-", "_") else "_"
            for c in filename
        )
        # Collapse multiple underscores
        while "__" in safe:
            safe = safe.replace("__", "_")
        return safe.strip("_") or "document.pdf"


# ── Module-level singleton ────────────────────────────────────────────────────
# Services import this instance directly — no need to instantiate.
pdf_service = PDFService()


# ── File-to-text converters ───────────────────────────────────────────────────

def _extract_text_from_file(file_bytes: bytes, suffix: str, filename: str) -> str:
    """
    Extract plain text from a non-PDF file.
    Supports: .docx, .doc, .txt, .xlsx, .xls, .csv

    Returns extracted text as a single string.
    Raises on unrecoverable errors.
    """
    import io as _io

    # ── Plain text ────────────────────────────────────────────────────────────
    if suffix == ".txt":
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    # ── CSV ───────────────────────────────────────────────────────────────────
    if suffix == ".csv":
        import csv as _csv
        text_io = _io.StringIO(file_bytes.decode("utf-8-sig", errors="replace"))
        reader  = _csv.reader(text_io)
        rows    = list(reader)
        if not rows:
            return ""
        # Build readable text: header as labels, rows as key:value lines
        header = rows[0]
        lines  = ["\t".join(header)]
        for row in rows[1:]:
            pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
            lines.append("  |  ".join(pairs))
        return "\n".join(lines)

    # ── DOCX ──────────────────────────────────────────────────────────────────
    if suffix in (".docx",):
        try:
            from docx import Document as _DocxDoc
            docx    = _DocxDoc(_io.BytesIO(file_bytes))
            parts   = [p.text for p in docx.paragraphs if p.text.strip()]
            # Also extract tables
            for table in docx.tables:
                for row in table.rows:
                    row_text = "  |  ".join(
                        c.text.strip() for c in row.cells if c.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n\n".join(parts)
        except ImportError:
            raise RuntimeError("python-docx not installed. Add it to requirements.txt.")

    # ── DOC (legacy Word) ─────────────────────────────────────────────────────
    if suffix == ".doc":
        # Try antiword-style extraction via python-docx (sometimes works)
        # Otherwise fall back to raw text extraction
        try:
            from docx import Document as _DocxDoc
            docx  = _DocxDoc(_io.BytesIO(file_bytes))
            parts = [p.text for p in docx.paragraphs if p.text.strip()]
            return "\n\n".join(parts)
        except Exception:
            # Raw fallback: extract printable ASCII from binary
            raw = file_bytes.decode("latin-1", errors="replace")
            return "\n".join(
                line for line in raw.splitlines()
                if len(line.strip()) > 20 and line.strip().isprintable()
            )

    # ── XLSX / XLS ────────────────────────────────────────────────────────────
    if suffix in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb    = openpyxl.load_workbook(_io.BytesIO(file_bytes), data_only=True)
            parts : list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(v) if v is not None else "" for v in row]
                    non_empty = [v for v in row_vals if v.strip()]
                    if non_empty:
                        parts.append("  |  ".join(non_empty))
            return "\n".join(parts)
        except ImportError:
            raise RuntimeError("openpyxl not installed. Add it to requirements.txt.")