"""
export_service.py — Export PDF metadata to XLSX / DOCX / CSV / JSON.

Two export templates:
  1. Journal Articles  — matches For_Metadata.xlsx (16 cols incl. name original)
  2. PhD Theses        — matches For_PhD_theses.xlsx (9 DC-style cols)

Field priority for every column:
  1. Schema field set by extraction_service
  2. Detected section content (e.g. abstract from sections list)
  3. Regex fallback from full_text

Fixes applied over previous version:
  - _fallback_authors() does NOT split on commas (preserves "Last, First")
  - _build_citation() handles hyphenated first names correctly
  - Error sentinel rows shown in XLSX with red styling
  - _extract_volume / _extract_issue split into dedicated functions
  - PhD theses export added (author, date, dc.description, abstract,
    citation, publisher, dc.subject, dc.title, dc.type)
  - Pydantic v2 compat: all getattr(model, field) calls use
    `getattr(m, field, None) or ""` to avoid AttributeError on undefined fields
  - _format_thesis_author: handles names with multiple commas (e.g. "Smith, Jr.")
  - _collect_rows / _collect_thesis_rows: m.page_count guarded with getattr
  - _collect_rows: year field uses getattr with None fallback
  - _build_citation inner lambda renamed to avoid shadowing outer variable `m`
"""

from __future__ import annotations

import csv
import io
import json
import re
from datetime import datetime
from html import unescape

from app.services.pdf_service import pdf_service
from app.utils.logger import get_logger

logger = get_logger(__name__)

# ── Column templates ──────────────────────────────────────────────────────────

# Matches For_Metadata.xlsx exactly (+ "name original" prepended)
JOURNAL_COLUMNS = [
    "name original",  # PDF filename
    "authors",        # separated by ||
    "editor",         # journal editor
    "date",           # publication year / full date
    "page no",        # page range e.g. 7-14 or e398-e404
    "abstract",       # full abstract
    "sponsor",        # funding / sponsor statement
    "citation",       # formatted citation
    "doi",            # DOI
    "issn",           # ISSN
    "publisher",      # publisher name
    "keywords",       # separated by ||
    "title",          # paper title
    "type",           # article type
    "issue",          # journal issue
    "volume",         # journal volume
]

# Matches For_PhD_theses.xlsx exactly
THESIS_COLUMNS = [
    "author",          # single author (Last, First.)
    "date",            # publication year
    "dc.description",  # physical description e.g. "xvi, 172p."
    "abstract",        # full abstract
    "citation",        # formatted citation
    "publisher",       # department + university
    "dc.subject",      # keywords separated by ||
    "dc.title",        # full thesis title
    "dc.type",         # always "Thesis"
]

# Default alias used by callers
XLSX_COLUMNS = JOURNAL_COLUMNS


class ExportService:

    # ── Public API ────────────────────────────────────────────────────────────

    def export_xlsx(
        self,
        doc_ids  : list[str],
        filename : str = "pdf_metadata_export.xlsx",
        template : str = "journal",   # "journal" | "thesis"
    ) -> tuple[bytes, str]:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError("pip install openpyxl")

        columns = THESIS_COLUMNS if template == "thesis" else JOURNAL_COLUMNS
        rows    = (
            self._collect_thesis_rows(doc_ids)
            if template == "thesis"
            else self._collect_rows(doc_ids)
        )

        wb = Workbook()
        ws = wb.active
        ws.title = "Metadata"

        header_font  = Font(name="Arial", bold=True, color="FFFFFF", size=11)
        header_fill  = PatternFill("solid", start_color="1A1A1A")
        accent_fill  = PatternFill("solid", start_color="BF3A14")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin         = Side(style="thin", color="DDDDDD")
        border       = Border(left=thin, right=thin, top=thin, bottom=thin)

        ws.row_dimensions[1].height = 32
        for col_i, col_name in enumerate(columns, start=1):
            label          = col_name.replace("dc.", "").replace("_", " ").title()
            cell           = ws.cell(row=1, column=col_i, value=label)
            cell.font      = header_font
            cell.fill      = accent_fill if col_i == 1 else header_fill
            cell.alignment = header_align
            cell.border    = border

        row_fill_even = PatternFill("solid", start_color="F8F5F0")
        row_fill_odd  = PatternFill("solid", start_color="FFFFFF")
        error_fill    = PatternFill("solid", start_color="FFF0F0")
        data_font     = Font(name="Arial", size=10)
        error_font    = Font(name="Arial", size=10, color="CC0000", italic=True)
        data_align    = Alignment(vertical="top", wrap_text=True)

        for row_i, row in enumerate(rows, start=2):
            is_error = row.get("_error", False)
            fill     = error_fill if is_error else (
                row_fill_even if row_i % 2 == 0 else row_fill_odd
            )
            font = error_font if is_error else data_font
            ws.row_dimensions[row_i].height = 60
            for col_i, col_name in enumerate(columns, start=1):
                cell           = ws.cell(row=row_i, column=col_i, value=row.get(col_name, ""))
                cell.font      = font
                cell.fill      = fill
                cell.alignment = data_align
                cell.border    = border

        # Column widths
        journal_widths = {
            "name original": 28, "authors": 36, "editor": 22,
            "date": 12, "page no": 12, "abstract": 60,
            "sponsor": 30, "citation": 45, "doi": 32,
            "issn": 14, "publisher": 24, "keywords": 32,
            "title": 45, "type": 20, "issue": 10, "volume": 10,
        }
        thesis_widths = {
            "author": 28, "date": 12, "dc.description": 18,
            "abstract": 65, "citation": 50, "publisher": 35,
            "dc.subject": 30, "dc.title": 50, "dc.type": 12,
        }
        widths = thesis_widths if template == "thesis" else journal_widths
        for col_i, col_name in enumerate(columns, start=1):
            ws.column_dimensions[get_column_letter(col_i)].width = widths.get(col_name, 16)

        ws.freeze_panes = "A2"

        # Summary sheet
        ok_count  = sum(1 for r in rows if not r.get("_error"))
        err_count = len(rows) - ok_count
        ws2       = wb.create_sheet("Summary")
        ws2["A1"] = "Export Summary"
        ws2["A1"].font = Font(name="Arial", bold=True, size=14)
        summary_rows = [
            ("Template",        "PhD Theses" if template == "thesis" else "Journal Articles"),
            ("Total Documents", len(rows)),
            ("Exported OK",     ok_count),
            ("Errors",          err_count),
            ("Exported At",     datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
            ("Columns",         ", ".join(columns)),
        ]
        for r_i, (label, value) in enumerate(summary_rows, start=3):
            ws2.cell(row=r_i, column=1, value=label).font = Font(name="Arial", bold=True, size=11)
            ws2.cell(row=r_i, column=2, value=value)

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        logger.info("XLSX export (%s) — %d ok, %d errors", template, ok_count, err_count)
        return buf.read(), filename

    def export_csv(
        self,
        doc_ids  : list[str],
        filename : str = "pdf_metadata_export.csv",
        template : str = "journal",
    ) -> tuple[bytes, str]:
        columns = THESIS_COLUMNS if template == "thesis" else JOURNAL_COLUMNS
        rows    = (
            self._collect_thesis_rows(doc_ids)
            if template == "thesis"
            else self._collect_rows(doc_ids)
        )
        buf    = io.StringIO()
        writer = csv.DictWriter(
            buf, fieldnames=columns,
            extrasaction="ignore", lineterminator="\n",
        )
        writer.writeheader()
        writer.writerows(rows)
        ok_count = sum(1 for r in rows if not r.get("_error"))
        logger.info("CSV export (%s) — %d ok, %d errors", template, ok_count, len(rows) - ok_count)
        return buf.getvalue().encode("utf-8-sig"), filename

    def export_docx(
        self,
        doc_ids  : list[str],
        filename : str = "pdf_research_report.docx",
        template : str = "journal",
    ) -> tuple[bytes, str]:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches  # noqa: F401
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("pip install python-docx")

        doc  = Document()
        rows = [
            r for r in (
                self._collect_thesis_rows(doc_ids)
                if template == "thesis"
                else self._collect_rows(doc_ids)
            )
            if not r.get("_error")
        ]

        report_title = "PhD Theses Report" if template == "thesis" else "PDF Research Analysis Report"
        title_para   = doc.add_heading(report_title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title_para.runs:
            title_para.runs[0].font.color.rgb = RGBColor(0xBF, 0x3A, 0x14)

        sub = doc.add_paragraph(
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
            f"  ·  Documents: {len(rows)}"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].font.size = Pt(10)
            sub.runs[0].font.color.rgb = RGBColor(0x85, 0x7F, 0x76)

        doc.add_page_break()

        for i, row in enumerate(rows, start=1):
            title_field = "dc.title" if template == "thesis" else "title"
            entry_title = row.get(title_field) or row.get("_filename", f"Document {i}")
            h = doc.add_heading(f"{i}. {entry_title[:120]}", level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0xBF, 0x3A, 0x14)

            if template == "thesis":
                meta_fields = [
                    ("Author",       row.get("author",         "") or "—"),
                    ("Date",         row.get("date",           "") or "—"),
                    ("Publisher",    row.get("publisher",      "") or "—"),
                    ("Description",  row.get("dc.description", "") or "—"),
                    ("Subject",      row.get("dc.subject",     "") or "—"),
                    ("Type",         row.get("dc.type",        "") or "—"),
                ]
            else:
                vol_issue = _format_vol_issue(row.get("volume", ""), row.get("issue", ""))
                meta_fields = [
                    ("Authors",     row.get("authors",   "") or "—"),
                    ("Editor",      row.get("editor",    "") or "—"),
                    ("Journal",     row.get("_journal",  "") or "—"),
                    ("Vol / Issue", vol_issue or "—"),
                    ("Date",        row.get("date",      "") or "—"),
                    ("Pages",       row.get("page no",   "") or "—"),
                    ("Publisher",   row.get("publisher", "") or "—"),
                    ("DOI",         row.get("doi",       "") or "—"),
                    ("ISSN",        row.get("issn",      "") or "—"),
                    ("Keywords",    row.get("keywords",  "") or "—"),
                    ("Type",        row.get("type",      "") or "—"),
                    ("Sponsor",     row.get("sponsor",   "") or "—"),
                ]

            table = doc.add_table(rows=len(meta_fields), cols=2)
            table.style = "Table Grid"
            for r_i, (label, value) in enumerate(meta_fields):
                lc, vc = table.cell(r_i, 0), table.cell(r_i, 1)
                lc.text = label
                vc.text = str(value)
                for run in lc.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
                for run in vc.paragraphs[0].runs:
                    run.font.size = Pt(9)

            doc.add_paragraph()

            abstract = row.get("abstract", "")
            if abstract:
                doc.add_heading("Abstract", level=2)
                p = doc.add_paragraph(abstract[:2000])
                for run in p.runs:
                    run.font.size   = Pt(10)
                    run.font.italic = True

            citation = row.get("citation", "")
            if citation:
                doc.add_heading("Citation", level=2)
                p = doc.add_paragraph(citation)
                for run in p.runs:
                    run.font.size = Pt(9)

            if i < len(rows):
                doc.add_page_break()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        logger.info("DOCX export (%s) — %d documents", template, len(rows))
        return buf.read(), filename

    def export_json(
        self,
        doc_ids         : list[str],
        filename        : str  = "pdf_metadata_export.json",
        template        : str  = "journal",
        include_internal: bool = False,
    ) -> tuple[bytes, str]:
        rows = (
            self._collect_thesis_rows(doc_ids)
            if template == "thesis"
            else self._collect_rows(doc_ids)
        )
        if not include_internal:
            rows = [{k: v for k, v in r.items() if not k.startswith("_")} for r in rows]

        ok_count = sum(1 for r in rows if not r.get("_error", False))
        data = {
            "exported_at"  : datetime.utcnow().isoformat(),
            "template"     : template,
            "total"        : len(rows),
            "total_ok"     : ok_count,
            "total_errors" : len(rows) - ok_count,
            "documents"    : rows,
        }
        logger.info("JSON export (%s) — %d ok, %d errors", template, ok_count, len(rows) - ok_count)
        return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8"), filename

    # ── Journal Articles data collection ─────────────────────────────────────

    def _collect_rows(self, doc_ids: list[str]) -> list[dict]:
        seen: set[str] = set()
        rows: list[dict] = []

        for doc_id in doc_ids:
            if doc_id in seen:
                continue
            seen.add(doc_id)

            try:
                doc = pdf_service.load_document(doc_id)
                if not doc:
                    rows.append(_error_row(doc_id, "Document not found"))
                    continue

                m      = doc.metadata
                first3 = "\n".join((doc.full_text or "").split("\n\n")[:60])
                full   = doc.full_text or ""

                title        = _clean(m.title or "") or _clean(doc.filename.replace(".pdf", ""))
                authors_list = _dedupe_authors(
                    [_clean(a) for a in (m.authors or []) if a.strip()]
                    or _fallback_authors(full, title)
                )
                authors      = " || ".join(authors_list)

                # FIX: use getattr with None default for optional Pydantic v2 fields
                editor    = _clean(getattr(m, "editor",  None) or "") or _extract_editor_fb(first3)
                year      = _clean(getattr(m, "year",    None) or "")
                date      = year or _parse_date(getattr(m, "created_at", None) or "")
                pages_raw = _clean(getattr(m, "pages",   None) or "") or _extract_pages_fb(first3)
                # FIX: guard page_count — may not exist on all schema versions
                page_count_val = getattr(m, "page_count", None) or 0
                page_no   = pages_raw or (str(page_count_val) if page_count_val else "")

                abstract = ""
                for sec in (doc.sections or []):
                    if sec.section_type.value == "abstract":
                        abstract = sec.content[:2000].strip()
                        break
                if not abstract:
                    abstract = _clean(getattr(m, "abstract", None) or "")[:2000]
                if not abstract:
                    abstract = _extract_abstract_fb(full[:5000])

                sponsor   = _clean(getattr(m, "funding",       None) or "") or _extract_funding_fb(full[:8000])
                doi       = _clean(getattr(m, "doi",           None) or "") or _extract_doi(first3)
                issn      = _clean(getattr(m, "issn",          None) or "") or _extract_issn(first3)
                publisher = _clean(getattr(m, "publisher",     None) or "") or _extract_publisher(first3)
                journal   = _clean(getattr(m, "journal",       None) or "") or _extract_journal(first3)
                volume    = _clean(getattr(m, "volume",        None) or "") or _extract_volume(first3)
                issue     = _clean(getattr(m, "issue",         None) or "") or _extract_issue(first3)
                kws       = list(getattr(m, "keywords", None) or []) or _extract_keywords_list(full)
                keywords  = " || ".join(_clean(k) for k in kws if k.strip())
                art_type  = _clean(getattr(m, "article_type",  None) or "") or _extract_article_type_fb(first3)
                citation  = _build_citation(
                    title=title, authors=authors_list, date=date,
                    journal=journal, volume=volume, issue=issue,
                    pages=pages_raw, doi=doi, publisher=publisher,
                )

                rows.append({
                    "name original" : doc.filename,
                    "authors"       : authors,
                    "editor"        : editor,
                    "date"          : date,
                    "page no"       : page_no,
                    "abstract"      : abstract,
                    "sponsor"       : sponsor,
                    "citation"      : citation,
                    "doi"           : doi,
                    "issn"          : issn,
                    "publisher"     : publisher,
                    "keywords"      : keywords,
                    "title"         : title,
                    "type"          : art_type or "Research Article",
                    "issue"         : issue,
                    "volume"        : volume,
                    "_filename"     : doc.filename,
                    "_doc_id"       : doc.doc_id,
                    "_journal"      : journal,
                    "_error"        : False,
                })

            except Exception as e:
                logger.error("Export failed for %s: %s", doc_id, e, exc_info=True)
                rows.append(_error_row(doc_id, str(e)))

        return rows

    # ── PhD Theses data collection ────────────────────────────────────────────

    def _collect_thesis_rows(self, doc_ids: list[str]) -> list[dict]:
        """
        Maps extracted metadata to the For_PhD_theses.xlsx column format:
          author | date | dc.description | abstract | citation |
          publisher | dc.subject | dc.title | dc.type
        """
        seen: set[str] = set()
        rows: list[dict] = []

        for doc_id in doc_ids:
            if doc_id in seen:
                continue
            seen.add(doc_id)

            try:
                doc = pdf_service.load_document(doc_id)
                if not doc:
                    rows.append(_error_thesis_row(doc_id, "Document not found"))
                    continue

                m      = doc.metadata
                first3 = "\n".join((doc.full_text or "").split("\n\n")[:60])
                full   = doc.full_text or ""

                # dc.title
                dc_title = _clean(getattr(m, "title", None) or "") or _clean(doc.filename.replace(".pdf", ""))

                # author — thesis usually has one author
                # FIX: all getattr calls use None default for Pydantic v2 compat
                authors_list = _dedupe_authors(
                    [_clean(a) for a in (getattr(m, "authors", None) or []) if a.strip()]
                    or _fallback_authors(full, dc_title)
                )
                author = _format_thesis_author(authors_list[0]) if authors_list else ""

                # date
                year = _clean(getattr(m, "year", None) or "")
                date = year or _parse_date(getattr(m, "created_at", None) or "")
                date = date[:4] if date else ""

                # dc.description — physical description e.g. "xvi, 172p."
                dc_description = _extract_physical_description(full[:3000])

                # abstract
                abstract = ""
                for sec in (doc.sections or []):
                    if sec.section_type.value == "abstract":
                        abstract = sec.content[:2000].strip()
                        break
                if not abstract:
                    abstract = _clean(getattr(m, "abstract", None) or "")[:2000]
                if not abstract:
                    abstract = _extract_abstract_fb(full[:5000])

                # publisher — for theses: "Department, Faculty, University"
                publisher = _clean(getattr(m, "publisher", None) or "") or _extract_thesis_publisher(full[:3000])

                # dc.subject — keywords
                kws     = list(getattr(m, "keywords", None) or []) or _extract_keywords_list(full)
                dc_subj = " || ".join(_clean(k) for k in kws if k.strip())

                # citation — thesis format
                citation = _build_thesis_citation(
                    title     = dc_title,
                    author    = author,
                    date      = date,
                    publisher = publisher,
                )

                rows.append({
                    "author"        : author,
                    "date"          : date,
                    "dc.description": dc_description,
                    "abstract"      : abstract,
                    "citation"      : citation,
                    "publisher"     : publisher,
                    "dc.subject"    : dc_subj,
                    "dc.title"      : dc_title,
                    "dc.type"       : "Thesis",
                    "_filename"     : doc.filename,
                    "_doc_id"       : doc.doc_id,
                    "_error"        : False,
                })

            except Exception as e:
                logger.error("Thesis export failed for %s: %s", doc_id, e, exc_info=True)
                rows.append(_error_thesis_row(doc_id, str(e)))

        return rows


# ══════════════════════════════════════════════════════════════════════════════
# Helper functions
# ══════════════════════════════════════════════════════════════════════════════

def _error_row(doc_id: str, reason: str) -> dict:
    row = {col: "" for col in JOURNAL_COLUMNS}
    row.update({
        "name original": doc_id,
        "title":         f"[Export error: {reason}]",
        "_doc_id":       doc_id,
        "_error":        True,
    })
    return row


def _error_thesis_row(doc_id: str, reason: str) -> dict:
    row = {col: "" for col in THESIS_COLUMNS}
    row.update({
        "dc.title":  f"[Export error: {reason}]",
        "_doc_id":   doc_id,
        "_error":    True,
    })
    return row


def _clean(text: str) -> str:
    return unescape(str(text)).strip()


def _dedupe_authors(authors: list[str]) -> list[str]:
    seen:   set[str]  = set()
    result: list[str] = []
    for a in authors:
        key = a.lower().strip()
        if key and key not in seen:
            seen.add(key)
            result.append(a)
    return result


def _format_vol_issue(volume: str, issue: str) -> str:
    parts: list[str] = []
    if volume:
        parts.append(f"Vol {volume}")
    if issue:
        parts.append(f"No {issue}")
    return ", ".join(parts)


def _format_thesis_author(name: str) -> str:
    """
    Format a thesis author name as "Last, First Middle."
    e.g. "Ezekiel Oluwakayode Idowu" → "Idowu, Ezekiel Oluwakayode."

    FIX: handles names with multiple commas (e.g. "Smith, John, Jr.") by
    only treating the part before the FIRST comma as the last name indicator,
    i.e. if a comma is already present we treat it as already-formatted.
    """
    name = name.strip().rstrip(".")
    if not name:
        return name

    if "," in name:
        # Already "Last, First [suffix]" style — normalise trailing period only
        # but do NOT re-split on additional commas (e.g. "Jr.")
        return name.rstrip(".") + "."

    tokens = name.split()
    if not tokens:
        return name
    last = tokens[-1]
    rest = " ".join(tokens[:-1])
    if rest:
        return f"{last}, {rest}."
    return f"{last}."


def _extract_physical_description(text: str) -> str:
    """
    Extract physical description from thesis text.
    Looks for patterns like "xvi, 172p." or "xx, 300 pages"
    """
    m = re.search(
        r"\b([ivxlcdmIVXLCDM]+(?:,\s*\d+\s*(?:p(?:ages?|\.)?|leaves?)))\b",
        text[:2000],
    )
    if m:
        return m.group(1).strip()
    m = re.search(r"\b(\d{2,4})\s*(?:pages?|p\.)\b", text[:2000], re.IGNORECASE)
    if m:
        return f"{m.group(1)}p."
    return ""


def _extract_thesis_publisher(text: str) -> str:
    """
    Extract thesis publisher (department + university).
    Looks for "Department of X, Faculty of Y, University Z"
    """
    m = re.search(
        r"(Department\s+of\s+[^\n]{5,80})",
        text, re.IGNORECASE,
    )
    if m:
        dept  = m.group(1).strip().rstrip(".,")
        rest  = text[m.end():]
        uni_m = re.search(
            r"((?:Obafemi|University|Polytechnic|College)[^\n]{3,60})",
            rest[:200], re.IGNORECASE,
        )
        if uni_m:
            return f"{dept}, {uni_m.group(1).strip().rstrip('.,')}"
        return dept
    # Fallback: any university name
    m = re.search(
        r"((?:Obafemi\s+Awolowo|University\s+of|Federal\s+University)[^\n]{3,60})",
        text, re.IGNORECASE,
    )
    return m.group(1).strip().rstrip(".,") if m else ""


def _build_thesis_citation(
    title     : str,
    author    : str,
    date      : str,
    publisher : str,
) -> str:
    """
    Build a thesis citation:
    Last, F. (Year). Title. Department, University.
    """
    parts: list[str] = []

    # Convert "Idowu, Ezekiel Oluwakayode." → "Idowu, E. O."
    if author:
        name = author.rstrip(".")
        if "," in name:
            last, rest = name.split(",", 1)
            tokens   = rest.strip().split()
            initials = " ".join(f"{t[0].upper()}." for t in tokens if t)
            parts.append(f"{last.strip()}, {initials}")
        else:
            parts.append(name)

    year = date[:4] if date else ""
    parts.append(f"({year})." if re.match(r"^(19|20)\d{2}$", year) else "(n.d.).")

    if title:
        words = title.split()
        cased = []
        for i, word in enumerate(words):
            core = re.sub(r"[^A-Za-z]", "", word)
            if core and core == core.upper() and len(core) >= 2:
                cased.append(word)
            elif i == 0:
                cased.append(word[0].upper() + word[1:].lower())
            else:
                cased.append(word.lower())
        t = " ".join(cased)
        # FIX: renamed lambda arg to avoid shadowing outer variable `m`
        t = re.sub(r"(:\s+)([a-z])", lambda mo: mo.group(1) + mo.group(2).upper(), t)
        parts.append(f"{t}.")

    if publisher:
        parts.append(f"{publisher}.")

    return " ".join(parts)


def _parse_date(raw: str) -> str:
    if not raw:
        return ""
    raw = raw.strip()
    if re.match(r"^(19|20)\d{2}$", raw):
        return raw
    match = re.match(r"D:(\d{4})(\d{2})(\d{2})", raw)
    if match:
        return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
    raw_clean = re.sub(r"(\d+)(st|nd|rd|th)\b", r"\1", raw)
    for fmt in (
        "%d %B %Y", "%B %d %Y", "%d %b %Y", "%b %d %Y",
        "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%B %Y", "%b %Y",
    ):
        try:
            return datetime.strptime(raw_clean.strip(), fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    yr = re.search(r"\b(19|20)\d{2}\b", raw)
    return yr.group(0) if yr else raw[:10]


def _build_citation(
    title     : str,
    authors   : list[str],
    date      : str,
    journal   : str = "",
    volume    : str = "",
    issue     : str = "",
    pages     : str = "",
    doi       : str = "",
    publisher : str = "",
) -> str:
    parts: list[str] = []

    if authors:
        formatted: list[str] = []
        for a in authors:
            a = a.strip()
            if not a:
                continue
            if "," in a:
                formatted.append(a)
            else:
                tokens   = a.split()
                last     = tokens[-1]
                initials = []
                for tok in tokens[:-1]:
                    if tok:
                        # FIX: handle hyphenated first names e.g. "Jean-Paul"
                        first_char = tok.lstrip("-")[0] if tok.lstrip("-") else tok[0]
                        initials.append(f"{first_char.upper()}.")
                init_str = " ".join(initials)
                formatted.append(f"{last}, {init_str}" if init_str else last)
        parts.append(" and ".join(formatted))

    year = (date[:4] if date and len(date) >= 4 else date) or ""
    parts.append(f"({year})." if re.match(r"^(19|20)\d{2}$", year) else "(n.d.).")

    if title:
        words = title.split()
        cased: list[str] = []
        for i, word in enumerate(words):
            core = re.sub(r"[^A-Za-z]", "", word)
            if core and core == core.upper() and len(core) >= 2:
                cased.append(word)
            elif "-" in word and any(p[:1].isupper() for p in word.split("-") if p):
                cased.append(word)
            elif i == 0:
                cased.append(word[0].upper() + word[1:].lower())
            else:
                cased.append(word.lower())
        t = " ".join(cased)
        # FIX: renamed lambda arg to avoid shadowing outer variable `m`
        t = re.sub(r"(:\s+)([a-z])", lambda mo: mo.group(1) + mo.group(2).upper(), t)
        parts.append(f"{t}.")

    if journal:
        src = journal
        if volume:
            src += f", {volume}"
            if issue:
                src += f"({issue})"
        if pages:
            src += f":{pages}"
        src += "."
        parts.append(src)
    elif publisher:
        parts.append(f"{publisher}.")

    if doi and not journal:
        parts.append(doi if doi.startswith("http") else f"https://doi.org/{doi}")

    return " ".join(parts)


# ── Fallback extractors ───────────────────────────────────────────────────────

def _extract_editor_fb(text: str) -> str:
    m = re.search(
        r"(?:Edited\s+by|Guest\s+Editor|Section\s+Editor|"
        r"Editor[-\s]?in[-\s]?Chief|Handling\s+Editor|Academic\s+Editor)"
        r"[:\s]+([A-Z][^\n]{3,80})",
        text[:4000], re.IGNORECASE,
    )
    if m:
        raw = re.split(
            r"\s*[,;]\s*(?:PhD|MD|Dr|Prof|University|Institute)",
            m.group(1), flags=re.IGNORECASE,
        )[0]
        return raw.strip().rstrip(".,;")[:80]
    return ""


def _extract_pages_fb(text: str) -> str:
    m = re.search(r"\bpp?\.?\s*([eE]?\d{1,6}\s*[-–]\s*[eE]?\d{1,6})\b", text[:3000], re.IGNORECASE)
    if m:
        return m.group(1).replace(" ", "").replace("–", "-")
    m = re.search(r"\bVol[^\n]{1,30},\s*([eE]?\d{1,6}[-–][eE]?\d{1,6})\b", text[:3000], re.IGNORECASE)
    if m:
        return m.group(1).replace("–", "-")
    m = re.search(r"\b\d+:\s*([eE]?\d{1,6}[-–][eE]?\d{1,6})\b", text[:3000])
    if m:
        return m.group(1).replace("–", "-")
    return ""


def _extract_abstract_fb(text: str) -> str:
    m = re.search(
        r"\b(?:Abstract|Summary)\b\s*[:—]?\s*\n?([\s\S]{80,2000}?)"
        r"(?=\n\s*\n\s*(?:Keywords?|Introduction|Background|1\.|$))",
        text, re.IGNORECASE,
    )
    if m:
        return _clean(m.group(1))[:2000]
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if len(p.strip()) > 120]
    for para in paras[1:5]:
        if len(para) > 120:
            return _clean(para[:2000])
    return ""


def _extract_funding_fb(text: str) -> str:
    for pat in [
        r"(?:Funding|Funding\s+source|Financial\s+support|Grant)[:\s]+([^\n]{10,300})",
        r"(?:supported\s+by|funded\s+by|sponsored\s+by)\s+([^\n]{10,200})",
        r"This\s+(?:study|work|research)\s+was\s+(?:supported|funded)\s+by\s+([^\n]{10,200})",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1).strip()[:300]
    return ""


def _extract_article_type_fb(text: str) -> str:
    for pat in [
        r"\b(Systematic\s+Review(?:\s+and\s+Meta[-\s]?Analysis)?)\b",
        r"\b(Meta[-\s]?Analysis)\b",
        r"\b(Randomized\s+Controlled\s+Trial|RCT)\b",
        r"\b(Clinical\s+Trial)\b",
        r"\b(Case\s+Report)\b",
        r"\b(Review\s+Article|Literature\s+Review|Narrative\s+Review|Scoping\s+Review|Review\s+Paper)\b",
        r"\b(Original\s+(?:Research|Article|Paper))\b",
        r"\b(Research\s+Article|Research\s+Paper)\b",
        r"\b(Short\s+(?:Communication|Report|Note))\b",
        r"\b(Letter\s+to\s+the\s+Editor)\b",
        r"\b(Technical\s+(?:Note|Report))\b",
    ]:
        m = re.search(pat, text[:3000], re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return "Research Article"


def _fallback_authors(text: str, title: str) -> list[str]:
    lines       = [line.strip() for line in text[:3000].split("\n") if line.strip()]
    title_found = False
    candidates  : list[str] = []
    title_low   = title[:30].lower() if title else ""
    _STOP = re.compile(
        r"\b(university|department|institute|college|abstract|"
        r"introduction|background|email|@|http|received|accepted|"
        r"copyright|doi|keywords?|school\s+of|faculty)\b",
        re.IGNORECASE,
    )
    for line in lines:
        if not title_found:
            if title_low and title_low[:15] in line.lower():
                title_found = True
            continue
        if _STOP.search(line) or len(line) > 200:
            break
        line = re.sub(r"^[\d,*†‡§¶\s]+", "", line).strip()
        if not line:
            continue
        # FIX: split on " and " or ";" — NOT commas (names use "Last, First")
        parts = re.split(r"\s+and\s+|\s*;\s*", line, flags=re.IGNORECASE)
        found_here = 0
        for part in parts:
            part = part.strip()
            if re.match(r"^[A-Z\u00C0-\u024F]", part) and re.search(r"[a-z]", part) and 4 < len(part) < 70:
                candidates.append(part)
                found_here += 1
                if len(candidates) >= 10:
                    break
        if found_here == 0 and candidates:
            break
        if len(candidates) >= 10:
            break
    return candidates[:10]


def _extract_doi(text: str) -> str:
    flat = re.sub(r"\n\s*", " ", text[:5000])
    m = re.search(r"https?://(?:dx\.)?doi\.org/(10\.\d{4,9}/[^\s\"'<>\]\)]+)", flat)
    if m:
        return m.group(1).rstrip(".,;)]")
    m = re.search(r"\bdoi\s*:?\s*(10\.\d{4,9}/[^\s\"'<>\]\)]+)", flat, re.IGNORECASE)
    if m:
        return m.group(1).rstrip(".,;)]")
    m = re.search(r"\b(10\.\d{4,9}/[^\s\"'<>\]\)]{3,})", flat)
    return m.group(1).rstrip(".,;)]") if m else ""


def _extract_issn(text: str) -> str:
    m = re.search(r"\b[EP]-?ISSN[:\s]*(\d{4}-\d{3}[\dXx])\b", text[:4000], re.IGNORECASE)
    if m:
        return m.group(1)
    m = re.search(r"\bISSN[:\s]*(\d{4}-\d{3}[\dXx])\b", text[:4000], re.IGNORECASE)
    if m:
        return m.group(1)
    m = re.search(
        r"(?:journal|issn|copyright|print|online)[^\n]*\b(\d{4}-\d{3}[\dXx])\b",
        text[:4000], re.IGNORECASE,
    )
    if m:
        return m.group(1)
    m = re.search(r"\b(\d{4}-\d{3}[\dXx])\b", text[:2000])
    return m.group(1) if m else ""


def _extract_publisher(text: str) -> str:
    known = [
        "Wolters Kluwer","Lippincott Williams","Elsevier","Cell Press",
        "Springer","Springer Nature","Nature Publishing","Wiley","Wiley-Blackwell",
        "Taylor & Francis","Taylor and Francis","BMJ Publishing","Sage Publications","SAGE",
        "Oxford University Press","Cambridge University Press","PLOS","BioMed Central","BMC",
        "Frontiers Media","MDPI","Hindawi","Dove Medical Press",
        "American Chemical Society","Royal Society of Chemistry","IEEE","ACM","Karger","Thieme",
        "African Journals Online","AJOL",
    ]
    chunk = text[:5000]
    m = re.search(
        r"(?:Published\s+by|Publisher\s*:|©\s*\d{4}\s+)([A-Z][^\n]{3,70})",
        chunk, re.IGNORECASE,
    )
    if m:
        pub = re.split(r"\s*(?:Inc\.|Ltd\.?|All rights|Copyright|\d{4})", m.group(1))[0]
        return pub.strip()[:80]
    cl = chunk.lower()
    for pub in sorted(known, key=len, reverse=True):
        if pub.lower() in cl:
            return pub
    return ""


def _extract_journal(text: str) -> str:
    for pat in [
        r"(?:published\s+in|journal\s*:)\s*([A-Z][^\n]{5,100})",
        r"((?:International|European|American|British|African|Asian|Nigerian|Indian|Chinese|Korean|Canadian|Australian)\s+(?:Journal|Review|Annals|Archives|Bulletin|Proceedings|Transactions|Letters|Reports)\s+(?:of|for|on|in)\s+[A-Z][^\n]{3,70})",
        r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,6}\s+Journal[^\n]{0,40})",
        r"(American Journal of [^\n]{5,60})",
        r"(British Journal of [^\n]{5,60})",
        r"(Journal of [A-Z][^\n]{5,60})",
    ]:
        m = re.search(pat, text[:6000], re.IGNORECASE)
        if m:
            j = m.group(1).strip()
            j = re.split(r"\s+(?:\d{4}\b|\bVol|\bNo\b|\bIssue|\d+\s*[\(,])", j)[0]
            j = j.strip().rstrip(".,;:")
            if len(j) >= 8:
                return j[:120]
    return ""


def _extract_volume(text: str) -> str:
    m = re.search(r"\bVol(?:ume)?\.?\s*(\d+)", text[:5000], re.IGNORECASE)
    return m.group(1) if m else ""


def _extract_issue(text: str) -> str:
    m = re.search(r"\bVol(?:ume)?\.?\s*\d+\s*[\(,]\s*(\d+)\s*[\),]", text[:5000], re.IGNORECASE)
    if m:
        return m.group(1)
    m = re.search(r"\b(?:Issue|No\.?|Number|Num\.?)\s*\.?\s*(\d+)", text[:5000], re.IGNORECASE)
    return m.group(1) if m else ""


def _extract_keywords_list(text: str) -> list[str]:
    m = re.search(
        r"(?:Keywords?|Key\s+words?|Index\s+[Tt]erms?)\s*[:—]\s*([^\n]{10,600})",
        text[:12000], re.IGNORECASE,
    )
    if not m:
        return []
    kws = [k.strip().strip("•·-–—*") for k in re.split(r"[;,•·]", m.group(1))]
    return [k for k in kws if 2 < len(k) < 100 and not k.isdigit()][:20]


# ── Singleton ─────────────────────────────────────────────────────────────────
export_service = ExportService()